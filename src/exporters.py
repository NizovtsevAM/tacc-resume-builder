"""
Exporters — DOCX, HTML, JSON resume generation.
"""

from __future__ import annotations

import html as html_mod
import json
import logging
import os
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .config import Settings
from .docx_styles import get_docx_theme
from .models import ResumeProfile

logger = logging.getLogger("TACC Resume builder")


def export_resume_json(profile: ResumeProfile, output_path: str) -> None:
    """Export resume data to JSON."""
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(profile.to_dict(), f, indent=2, ensure_ascii=False)
    logger.info("Resume JSON exported to: %s", output_path)


def save_document(doc: Document, output_path: str) -> None:
    """Save the DOCX document."""
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    logger.info("Resume DOCX saved to: %s", output_path)


def export_resume_html(profile: ResumeProfile, settings: Settings, output_path: str) -> None:
    """Export resume data to HTML using a template."""
    templates_dir = os.path.join(os.path.dirname(__file__), "..", "templates")
    template_name = settings.resume_template if settings.resume_template else "modern"
    template_path = os.path.join(templates_dir, f"{template_name}.html")

    if not os.path.exists(template_path):
        logger.warning(
            "Template %s not found at %s, falling back to modern",
            template_name,
            template_path,
        )
        template_path = os.path.join(templates_dir, "modern.html")
        if not os.path.exists(template_path):
            logger.error("No templates found. Skipping HTML export.")
            return

    with open(template_path, "r", encoding="utf-8") as f:
        template = f.read()

    # Build contact string with hyperlinks
    contact_parts: list[str] = []
    if profile.email:
        contact_parts.append(
            f'<a href="mailto:{html_mod.escape(profile.email)}" style="color: #8892b0; text-decoration: none;">{html_mod.escape(profile.email)}</a>'
        )
    if profile.phone:
        contact_parts.append(html_mod.escape(profile.phone))
    if profile.location:
        contact_parts.append(html_mod.escape(profile.location))
    if profile.linkedin:
        linkedin_url = profile.linkedin
        if not linkedin_url.startswith("http"):
            linkedin_url = "https://" + linkedin_url
        contact_parts.append(
            f'<a href="{html_mod.escape(linkedin_url)}" target="_blank" rel="noopener" style="color: #8892b0; text-decoration: none;">{html_mod.escape(profile.linkedin)}</a>'
        )
    contacts = " | ".join(contact_parts)

    # Build skills section HTML
    skills_html = ""
    if profile.skills:
        skills_html = '<div class="section">\n  <h2>Technical Skills</h2>\n'
        for category in [
            "Backend",
            "Frontend",
            "Cloud",
            "DevOps",
            "Databases",
            "Testing",
            "Data Engineering",
            "Architecture",
        ]:
            if category in profile.skills:
                skills_html += (
                    f'  <div class="skill-category">\n'
                    f'    <div class="cat-name">{html_mod.escape(category)}</div>\n'
                    f'    <div class="skills-grid">\n'
                )
                for skill in sorted(profile.skills[category]):
                    skills_html += (
                        f'      <span class="skill-tag">{html_mod.escape(skill)}</span>\n'
                    )
                skills_html += "    </div>\n  </div>\n"
        for category in sorted(profile.skills.keys()):
            if category not in [
                "Backend",
                "Frontend",
                "Cloud",
                "DevOps",
                "Databases",
                "Testing",
                "Data Engineering",
                "Architecture",
            ]:
                skills_html += (
                    f'  <div class="skill-category">\n'
                    f'    <div class="cat-name">{html_mod.escape(category)}</div>\n'
                    f'    <div class="skills-grid">\n'
                )
                for skill in sorted(profile.skills[category]):
                    skills_html += (
                        f'      <span class="skill-tag">{html_mod.escape(skill)}</span>\n'
                    )
                skills_html += "    </div>\n  </div>\n"
        skills_html += "</div>\n"

    # Build projects section HTML
    projects_html = ""
    for project in profile.projects:
        projects_html += '<div class="project">\n'
        projects_html += f'  <div class="company">{html_mod.escape(project.customer)}</div>\n'
        if project.role:
            projects_html += f'  <div class="role">{html_mod.escape(project.role)}</div>\n'
        projects_html += (
            f'  <div class="dates">'
            f"{project.start_date.strftime('%b %Y')} — {project.end_date.strftime('%b %Y')}"
            f"  |  {project.days} days"
            f"</div>\n"
        )
        projects_html += '  <div class="details">\n'
        if project.responsibilities:
            projects_html += "    <h3>Responsibilities</h3>\n    <ul>\n"
            for resp in project.responsibilities:
                projects_html += f"      <li>{html_mod.escape(resp)}</li>\n"
            projects_html += "    </ul>\n"
        if project.achievements:
            projects_html += "    <h3>Achievements</h3>\n    <ul>\n"
            for ach in project.achievements:
                projects_html += f"      <li>{html_mod.escape(ach)}</li>\n"
            projects_html += "    </ul>\n"
        projects_html += "  </div>\n"
        if project.technologies:
            projects_html += (
                f'  <div class="techs">'
                f"<strong>Technologies:</strong> {html_mod.escape(', '.join(sorted(set(project.technologies))))}"
                f"</div>\n"
            )
        projects_html += "</div>\n"

    # Build education section
    education_html = ""
    if profile.education:
        education_html = (
            '<div class="section">\n'
            f"  <h2>Education</h2>\n"
            f"  <p>{html_mod.escape(profile.education)}</p>\n"
            "</div>\n"
        )

    # Replace placeholders
    html_content = template
    html_content = html_content.replace(
        "{{ NAME }}", html_mod.escape(f"{profile.first_name} {profile.last_name}")
    )
    html_content = html_content.replace("{{ PROFESSION }}", html_mod.escape(profile.profession))
    html_content = html_content.replace("{{ CONTACTS }}", contacts)
    html_content = html_content.replace("{{ SUMMARY }}", html_mod.escape(profile.summary))
    html_content = html_content.replace("{{ SKILLS_SECTION }}", skills_html)
    html_content = html_content.replace("{{ PROJECTS }}", projects_html)
    html_content = html_content.replace("{{ EDUCATION_SECTION }}", education_html)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html_content)
    logger.info("Resume HTML exported to: %s", output_path)


class ResumeDocumentGenerator:
    """Generates ATS-optimized professional DOCX resumes."""

    def __init__(self, settings: Settings) -> None:
        self.settings = settings
        self._theme = get_docx_theme(settings.resume_template)
        logger.info("Using DOCX template: %s", self._theme.name)

    def generate(self, profile: ResumeProfile) -> Document:
        """Create a professional DOCX document from the resume profile."""
        theme = self._theme
        doc = Document()

        section = doc.sections[0]
        section.top_margin = Inches(theme.margin_top)
        section.bottom_margin = Inches(theme.margin_bottom)
        section.left_margin = Inches(theme.margin_left)
        section.right_margin = Inches(theme.margin_right)

        style = doc.styles["Normal"]
        style.font.name = theme.font_name
        style.font.size = Pt(theme.body_size)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.line_spacing = 1.08

        self._add_header(doc, profile)

        self._add_section_heading(doc, "PROFESSIONAL SUMMARY")
        p = doc.add_paragraph(profile.summary)
        p.style.font.name = theme.font_name
        p.paragraph_format.space_after = Pt(6)

        if profile.skills and theme.show_core_competencies:
            self._add_section_heading(doc, "CORE COMPETENCIES")
            all_skills: list[str] = []
            for cat in [
                "Backend",
                "Frontend",
                "Cloud",
                "DevOps",
                "Databases",
                "Testing",
                "Data Engineering",
                "Architecture",
            ]:
                if cat in profile.skills:
                    all_skills.extend(profile.skills[cat])
            for cat, skills in profile.skills.items():
                if cat not in [
                    "Backend",
                    "Frontend",
                    "Cloud",
                    "DevOps",
                    "Databases",
                    "Testing",
                    "Data Engineering",
                    "Architecture",
                    "Other",
                ]:
                    all_skills.extend(skills)

            seen = set()
            unique_skills: list[str] = []
            for s in all_skills:
                if s.lower() not in seen:
                    unique_skills.append(s)
                    seen.add(s.lower())

            p = doc.add_paragraph(", ".join(unique_skills[:20]))
            p.style.font.name = theme.font_name
            p.paragraph_format.space_after = Pt(6)

        if profile.skills:
            self._add_section_heading(doc, "TECHNICAL SKILLS")
            for category in [
                "Backend",
                "Frontend",
                "Cloud",
                "DevOps",
                "Databases",
                "Testing",
                "Data Engineering",
                "Architecture",
            ]:
                if category in profile.skills:
                    self._add_skill_line(doc, category, profile.skills[category])
            for category in sorted(profile.skills.keys()):
                if category not in [
                    "Backend",
                    "Frontend",
                    "Cloud",
                    "DevOps",
                    "Databases",
                    "Testing",
                    "Data Engineering",
                    "Architecture",
                ]:
                    self._add_skill_line(doc, category, profile.skills[category])

        self._add_section_heading(doc, "PROFESSIONAL EXPERIENCE")
        for project in profile.projects:
            self._add_project_entry(doc, project)

        if profile.education:
            self._add_section_heading(doc, "EDUCATION")
            p = doc.add_paragraph(profile.education)
            p.style.font.name = theme.font_name

        return doc

    def _add_run(
        self,
        paragraph,
        text: str,
        *,
        size: float,
        bold: bool = False,
        italic: bool = False,
        color: RGBColor | None = None,
        font_name: str | None = None,
    ):
        theme = self._theme
        run = paragraph.add_run(text)
        run.font.size = Pt(size)
        run.font.name = font_name or theme.font_name
        run.bold = bold
        run.italic = italic
        if color is not None:
            run.font.color.rgb = color
        return run

    @staticmethod
    def _set_left_border(paragraph, color_hex: str, size: int = 18) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), str(size))
        left.set(qn("w:space"), "4")
        left.set(qn("w:color"), color_hex)
        p_bdr.append(left)
        p_pr.append(p_bdr)

    def _add_hyperlink(self, paragraph, text: str, url: str, font_size: float) -> None:
        theme = self._theme
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = paragraph._p.makeelement(qn("w:hyperlink"), {qn("r:id"): r_id})
        new_run = hyperlink.makeelement(qn("w:r"), {})
        r_pr = new_run.makeelement(qn("w:rPr"), {})
        c = r_pr.makeelement(qn("w:color"), {qn("w:val"): theme.link_color})
        r_pr.append(c)
        u = r_pr.makeelement(qn("w:u"), {qn("w:val"): "single"})
        r_pr.append(u)
        sz = r_pr.makeelement(qn("w:sz"), {qn("w:val"): str(int(font_size * 2))})
        r_pr.append(sz)
        fonts = r_pr.makeelement(
            qn("w:rFonts"),
            {
                qn("w:ascii"): theme.contact_font_name,
                qn("w:hAnsi"): theme.contact_font_name,
            },
        )
        r_pr.append(fonts)
        new_run.append(r_pr)
        t = new_run.makeelement(qn("w:t"), {})
        t.text = text
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    def _rule_color(self) -> RGBColor:
        theme = self._theme
        if theme.name == "creative":
            return theme.accent_color
        return RGBColor(0x95, 0xA5, 0xA6)

    def _rule_char(self) -> str:
        return "═" if self._theme.name == "classic" else "─"

    def _add_header(self, doc: Document, profile: ResumeProfile) -> None:
        theme = self._theme

        p = doc.add_paragraph()
        p.alignment = theme.header_align
        p.paragraph_format.space_after = Pt(0)
        self._add_run(
            p,
            f"{profile.first_name} {profile.last_name}",
            size=theme.name_size,
            bold=True,
            color=theme.name_color,
        )

        p = doc.add_paragraph()
        p.alignment = theme.header_align
        p.paragraph_format.space_after = Pt(2)
        self._add_run(
            p,
            profile.profession,
            size=theme.title_size,
            italic=theme.title_italic,
            color=theme.title_color,
        )

        contact_parts: list[tuple[str, str | None]] = []
        if profile.email:
            contact_parts.append((profile.email, f"mailto:{profile.email}"))
        if profile.phone:
            contact_parts.append((profile.phone, None))
        if profile.location:
            contact_parts.append((profile.location, None))
        if profile.linkedin:
            linkedin_url = profile.linkedin
            if not linkedin_url.startswith("http"):
                linkedin_url = "https://" + linkedin_url
            contact_parts.append((profile.linkedin, linkedin_url))

        if contact_parts:
            p = doc.add_paragraph()
            p.alignment = theme.header_align
            p.paragraph_format.space_after = Pt(4)
            for i, (text, url) in enumerate(contact_parts):
                if i > 0:
                    self._add_run(
                        p,
                        " | ",
                        size=theme.contact_size,
                        color=theme.contact_color,
                        font_name=theme.contact_font_name,
                    )
                if url:
                    self._add_hyperlink(p, text, url, theme.contact_size)
                else:
                    self._add_run(
                        p,
                        text,
                        size=theme.contact_size,
                        color=theme.contact_color,
                        font_name=theme.contact_font_name,
                    )

        if theme.header_rule:
            p = doc.add_paragraph()
            p.alignment = theme.header_align
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            self._add_run(
                p,
                self._rule_char() * theme.section_rule_width,
                size=6,
                color=self._rule_color(),
            )

    def _add_section_heading(self, doc: Document, text: str) -> None:
        theme = self._theme
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        if theme.header_align == WD_ALIGN_PARAGRAPH.LEFT:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        self._add_run(
            p,
            text.upper(),
            size=theme.section_heading_size,
            bold=True,
            color=theme.heading_color,
        )
        self._add_run(
            p,
            "\n" + self._rule_char() * theme.section_rule_width,
            size=5,
            color=self._rule_color(),
        )

    def _add_skill_line(self, doc: Document, category: str, skills: list[str]) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.keep_with_next = True

        self._add_run(p, f"{category}: ", size=10, bold=True)
        self._add_run(p, ", ".join(sorted(skills)), size=10)

    def _project_indent(self) -> Inches:
        return Inches(0.12) if self._theme.project_left_border else Inches(0)

    def _add_project_entry(self, doc: Document, project: Any) -> None:
        theme = self._theme
        indent = self._project_indent()

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.keep_with_next = True
        if theme.project_left_border:
            p.paragraph_format.left_indent = indent
            self._set_left_border(p, theme.accent_hex)

        self._add_run(
            p,
            project.customer,
            size=theme.company_size,
            bold=True,
            color=theme.heading_color,
        )

        if project.role:
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(0)
            p2.paragraph_format.keep_with_next = True
            if theme.project_left_border:
                p2.paragraph_format.left_indent = indent

            self._add_run(
                p2,
                project.role,
                size=theme.role_size,
                bold=theme.role_bold,
                italic=theme.role_italic,
                color=theme.role_color,
            )

        p3 = doc.add_paragraph()
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after = Pt(4)
        p3.paragraph_format.keep_with_next = True
        if theme.project_left_border:
            p3.paragraph_format.left_indent = indent

        date_font = theme.contact_font_name if theme.name == "creative" else theme.font_name
        self._add_run(
            p3,
            f"{project.start_date.strftime('%b %Y')} — {project.end_date.strftime('%b %Y')}"
            f"  |  {project.days} days",
            size=theme.date_size,
            color=theme.date_color,
            font_name=date_font,
        )

        if project.responsibilities:
            p4 = doc.add_paragraph()
            p4.paragraph_format.space_before = Pt(2)
            p4.paragraph_format.space_after = Pt(2)
            if theme.project_left_border:
                p4.paragraph_format.left_indent = indent
            self._add_run(p4, "Responsibilities:", size=10, bold=True)

            for resp in project.responsibilities:
                self._add_bullet(doc, resp, indent)

        if project.achievements:
            p5 = doc.add_paragraph()
            p5.paragraph_format.space_before = Pt(4)
            p5.paragraph_format.space_after = Pt(2)
            if theme.project_left_border:
                p5.paragraph_format.left_indent = indent
            self._add_run(p5, "Achievements:", size=10, bold=True)

            for ach in project.achievements:
                self._add_bullet(doc, ach, indent)

        if project.technologies:
            p6 = doc.add_paragraph()
            p6.paragraph_format.space_before = Pt(4)
            p6.paragraph_format.space_after = Pt(2)
            if theme.project_left_border:
                p6.paragraph_format.left_indent = indent
            self._add_run(p6, "Technologies: ", size=9.5, bold=True)
            tech_font = theme.contact_font_name if theme.name == "creative" else theme.font_name
            self._add_run(
                p6,
                ", ".join(sorted(set(project.technologies))),
                size=9.5,
                color=theme.tech_color,
                font_name=tech_font,
            )

        p_spacer = doc.add_paragraph()
        p_spacer.paragraph_format.space_before = Pt(2)
        p_spacer.paragraph_format.space_after = Pt(0)
        self._add_run(p_spacer, "", size=2)

    def _add_bullet(self, doc: Document, text: str, indent: Inches) -> None:
        theme = self._theme
        if theme.bullet_prefix:
            bullet = doc.add_paragraph()
            bullet.paragraph_format.space_after = Pt(1)
            bullet.paragraph_format.space_before = Pt(0)
            if theme.project_left_border:
                bullet.paragraph_format.left_indent = indent
            self._add_run(bullet, theme.bullet_prefix, size=10, color=theme.accent_color)
            self._add_run(bullet, text, size=10)
            return

        bullet = doc.add_paragraph(style="List Bullet")
        bullet.clear()
        self._add_run(bullet, text, size=10)
        bullet.paragraph_format.space_after = Pt(1)
        bullet.paragraph_format.space_before = Pt(0)
